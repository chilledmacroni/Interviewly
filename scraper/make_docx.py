from docx import Document

doc = Document()
doc.add_heading('John Doe', level=1)
doc.add_paragraph('Software Engineer')
doc.add_paragraph('Experienced backend engineer with 5+ years building distributed services in .NET and Python. Skilled in system design, microservices, and databases.')
doc.add_heading('Skills', level=2)
doc.add_paragraph('Python, C#, .NET, MongoDB, Docker, Kubernetes, REST, CI/CD')
doc.save('sample_resume.docx')
print('sample_resume.docx created')